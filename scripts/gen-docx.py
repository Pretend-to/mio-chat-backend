from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

doc = Document()

# ── Default style ──
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
style.paragraph_format.line_spacing = 1.35
style.paragraph_format.space_after = Pt(6)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# ── Heading styles ──
for level, (size, color_hex) in enumerate([
    (22, '1a1a2e'),
    (16, '16213e'),
    (13, '0f3460'),
    (12, '533483'),
], start=0):
    if level == 0:
        h = doc.styles['Title']
    else:
        h = doc.styles[f'Heading {level}']
    h.font.name = '微软雅黑'
    h.font.size = Pt(size)
    h.font.color.rgb = RGBColor(*[int(color_hex[i:i+2], 16) for i in (0, 2, 4)])
    h.font.bold = True
    h.paragraph_format.space_before = Pt(16 if level > 0 else 0)
    h.paragraph_format.space_after = Pt(8)
    h.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# ── Helper ──
def add_img_placeholder(doc, label):
    """Add an image placeholder box"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'[ 截图占位：{label} ]')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.italic = True
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

# ═══════════════════════════════════════
# Title
# ═══════════════════════════════════════
title = doc.add_heading('Python + AI 教学助手', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('基于蜜柚（MioChat）智能 Agent 平台的 Python 辅助教学工具')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x53, 0x34, 0x83)

doc.add_paragraph()

# ═══════════════════════════════════════
# 1. 背景
# ═══════════════════════════════════════
doc.add_heading('一、背景', level=1)

doc.add_paragraph(
    '蜜柚（MioChat）是一款企业级多协议 AI 对话平台与 Agent 操作系统，'
    '具备 Skills（能力包）、MCP（标准协议）、Plugins（原生插件）三位一体的扩展架构。'
    'Python + AI 教学助手是基于蜜柚平台构建的一项专项能力，'
    '面向 Python 编程教学场景，提供代码审查与沙盒执行两大核心功能。'
)

# ═══════════════════════════════════════
# 2. 工具定位
# ═══════════════════════════════════════
doc.add_heading('二、工具定位与适用场景', level=1)

doc.add_paragraph(
    '该工具定位为 Python 编程教学的辅助手段，适用于课后作业批改、课堂代码演示、'
    '实验报告验证、算法训练等场景。学生在对话界面提交 Python 代码后，'
    '工具自动完成静态审查并执行代码，即时反馈结果。'
)

doc.add_paragraph(
    '工具采用 24 小时在线模式，学生可随时使用，不受教师工作时间限制。'
    '所有代码在安全沙盒中执行，不触及宿主机系统，无安全风险。'
)

# ═══════════════════════════════════════
# 3. 核心功能
# ═══════════════════════════════════════
doc.add_heading('三、核心功能', level=1)

# 3.1
doc.add_heading('3.1 八维静态代码审查', level=2)

doc.add_paragraph(
    '在不执行代码的前提下，从八个维度对提交的 Python 代码进行全面分析，'
    '涵盖正确性、可读性、性能、安全性、PEP 8 风格、健壮性、模块化程度及文档完整性。'
    '审查报告以结构化方式呈现，每项问题附带具体说明与改进建议。'
)

add_img_placeholder(doc, '代码审查报告界面截图——展示八维审查结果与改进建议')

# 3.2
doc.add_heading('3.2 安全沙盒执行', level=2)

doc.add_paragraph(
    '内置 Python 3.10 沙盒运行环境，预装教学常用库（numpy、matplotlib、pandas、'
    'scipy、Pillow、jieba、wordcloud），无需额外安装依赖。'
    '代码提交后即刻执行，标准输出、错误信息和生成的图表均自动返回。'
)

add_img_placeholder(doc, '沙盒执行结果截图——展示代码输出与运行状态')

# 3.3
doc.add_heading('3.3 中文可视化自动适配', level=2)

doc.add_paragraph(
    '针对中文教学场景，工具内置中文字体配置能力（WenQuanYi Zen Hei、'
    'Noto Sans CJK SC 等）。学生使用 matplotlib、wordcloud、Pillow 等库绘制图表时，'
    '无需手动配置中文字体，系统自动完成字体注入，消除中文乱码问题。'
)

add_img_placeholder(doc, '中文图表效果截图——展示自动适配后的中文标题与标签')

# 3.4
doc.add_heading('3.4 审查-修正-验证闭环', level=2)

doc.add_paragraph(
    '工具在输出审查报告后，会自动生成修正后的完整代码并立即执行验证，'
    '形成"发现问题 → 给出方案 → 验证效果"的闭环流程。学生可基于执行结果'
    '进一步迭代优化。'
)

# ═══════════════════════════════════════
# 4. 教学应用示例
# ═══════════════════════════════════════
doc.add_heading('四、教学应用示例', level=1)

doc.add_heading('4.1 课后作业批改', level=2)
doc.add_paragraph(
    '学生提交排序算法、数据结构或函数实现类作业后，工具自动检查逻辑正确性、'
    '边界条件处理、代码风格和算法效率，并执行验证输出结果。'
    '教师可直接引用审查报告作为批改参考。'
)
add_img_placeholder(doc, '冒泡排序作业审查与执行结果截图')

doc.add_heading('4.2 数据分析可视化', level=2)
doc.add_paragraph(
    '学生完成数据处理与分析任务后，工具执行代码并直接生成可视化图表（折线图、'
    '柱状图、饼图、箱线图等），无需在本地搭建 Python 环境。'
)
add_img_placeholder(doc, '成绩数据分析可视化截图')

doc.add_heading('4.3 词云与文本分析', level=2)
doc.add_paragraph(
    '利用预装的 jieba 分词和 wordcloud 库，可支持中文文本处理、词频统计、'
    '词云生成等 NLP 教学任务。'
)
add_img_placeholder(doc, '中文词云生成效果截图')

# ═══════════════════════════════════════
# 5. 使用方法
# ═══════════════════════════════════════
doc.add_heading('五、使用方法', level=1)

doc.add_paragraph(
    '在蜜柚平台中与该教学助手对话，直接提交 Python 代码即可触发自动审查与执行。'
    '支持以下唤醒方式：'
)

items = [
    '提交 Python 代码段，工具自动识别并启动审查',
    '发送"帮我 Review 这段 Python 代码"等指令主动触发',
    '发送代码文件，工具读取后自动处理',
]
for item in items:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(item)

# ═══════════════════════════════════════
# 6. 价值总结
# ═══════════════════════════════════════
doc.add_heading('六、价值总结', level=1)

table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = '维度'
hdr[1].text = '说明'
for cell in hdr:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

items = [
    ('学习反馈效率', '从教师批改的"天级"压缩到机器审查的"秒级"'),
    ('代码质量意识', '八维审查帮助学生建立良好的编码习惯和规范意识'),
    ('环境门槛', '无需本地安装 Python 环境，浏览器即可完成编码验证'),
    ('中文支持', '内置中文字体与分词支持，消除中文教学环境障碍'),
    ('教学辅助', '审查报告可作为教师批改参考，减轻重复性劳动'),
]
for dim, val in items:
    row = table.add_row().cells
    row[0].text = dim
    row[1].text = val

doc.add_paragraph()
p = doc.add_paragraph(
    'Python + AI 教学助手基于蜜柚（MioChat）Agent 操作系统构建。'
    '该工具旨在辅助 Python 编程教学，而非替代教师的专业判断。'
)

# ═══════════════════════════════════════
# Footer
# ═══════════════════════════════════════
p = doc.add_paragraph()
run = p.add_run('—— 文档完 ——')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ═══════════════════════════════════════
# Save
# ═══════════════════════════════════════
output_path = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    'docs', 'python-ai-tutor-intro.docx'
)
doc.save(output_path)
print(f"Saved: {output_path}")
print(f"Size: {os.path.getsize(output_path)} bytes")
